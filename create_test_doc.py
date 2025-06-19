from docx import Document
import os

# 确保目录存在
if not os.path.exists('student_reports'):
    os.makedirs('student_reports')

# 创建新的 Word 文档
doc = Document()

# 添加标题
doc.add_heading('Test Student Report', 0)

# 添加段落
doc.add_paragraph('This is a test report content.')
doc.add_paragraph('This report contains the following sections:')

# 添加一些测试内容
sections = [
    '1. Introduction',
    'In this section, we introduce the project background and objectives.',

    '2. Methods',
    'This section describes our research methods and tools.',

    '3. Results',
    'Here we present our research findings and data analysis results.',

    '4. Discussion',
    'This section discusses the implications and potential impact of our findings.',

    '5. Conclusion',
    'Finally, we summarize the main findings and recommendations of the study.'
]

for section in sections:
    doc.add_paragraph(section)

# 保存文档
doc.save('student_reports/test_report.docx')
print("Test document created: student_reports/test_report.docx")
