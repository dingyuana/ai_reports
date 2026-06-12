"""
文件名: document_processor.py
作用: 实现文档处理器，用于处理不同类型的学生报告文件（PDF和Word）
实现路径:
    1. 定义DocumentProcessor抽象基类，规定文档处理的基本接口
    2. 实现PDFProcessor类，处理PDF格式的报告文件
    3. 实现WordProcessor类，处理Word格式的报告文件（.docx和.doc）
功能:
    - 从文档中提取文本内容
    - 在文档上添加评语和分数
    - 在文档上添加对号标注
使用方式:
    - 在grading_system.py中被实例化并使用
    - 根据文件扩展名选择相应的处理器
"""

import os
import io
from typing import List, Dict, Any
from abc import ABC, abstractmethod
import logging
from datetime import datetime
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import Color
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pdfplumber
try:
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False

# 配置日志记录器
logger = logging.getLogger(__name__)

# 注册中文字体（尝试多种字体作为备选）
CHINESE_FONT_NAME = 'Helvetica'  # 默认使用Helvetica
CHINESE_FONT_PATH = None

# 按优先级顺序尝试不同的字体（TTC和TTF都支持）
font_candidates = [
    ('ChineseFont', '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc'),
    ('ChineseFont', '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc'),
    ('ChineseFont', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'),
    ('ChineseFont', '/usr/share/fonts/truetype/freefont/FreeSans.ttf'),
    ('ChineseFont', '/usr/share/fonts/truetype/ubuntu/Ubuntu-R.ttf'),
]

for font_name, font_path in font_candidates:
    if os.path.exists(font_path):
        try:
            pdfmetrics.registerFont(TTFont(font_name, font_path))
            CHINESE_FONT_NAME = font_name
            CHINESE_FONT_PATH = font_path
            logger.info(f"成功注册中文字体: {font_path}")
            break
        except Exception as e:
            logger.warning(f"注册字体失败 {font_path}: {e}")
            continue

if CHINESE_FONT_NAME == 'Helvetica':
    logger.warning("未找到可用的中文字体，将使用默认字体Helvetica")


class DocumentProcessor(ABC):
    """文档处理抽象基类"""

    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """从文档中提取文本"""
        pass

    @abstractmethod
    def add_comments_and_score(self, file_path: str, comments: str, score: int, output_path: str, add_score: bool = True) -> None:
        """在文档上添加评语和分数
        
        Args:
            file_path: 输入文档路径
            comments: 要添加的评语
            score: 要添加的分数
            output_path: 输出文档路径
            add_score: 是否添加分数，默认为True
        """
        pass

    @abstractmethod
    def add_checkmarks(self, file_path: str, output_path: str) -> None:
        """在每页添加对号标注"""
        pass


class PDFProcessor(DocumentProcessor):
    """PDF文档处理实现"""

    def extract_text(self, file_path: str) -> str:
        """从PDF中提取文本"""
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text

    def add_comments_and_score(self, file_path: str, comments: str, score: int, output_path: str, add_score: bool = True) -> None:
        """在PDF上添加评语和分数，保留原报告内容
        
        Args:
            file_path: 输入PDF路径
            comments: 要添加的评语
            score: 要添加的分数
            output_path: 输出PDF路径
            add_score: 是否添加分数，默认为True
        """
        input_pdf = PdfReader(file_path)
        output_pdf = PdfWriter()
        
        # 处理第一页
        first_page = input_pdf.pages[0]
        
        # 只在需要时添加分数
        if add_score:
            # 创建分数标注
            score_packet = io.BytesIO()
            score_canvas = canvas.Canvas(score_packet, pagesize=letter)
            score_canvas.setFont("Helvetica-Bold", 48)  # 字号增加到48pt
            score_canvas.setFillColor(Color(1, 0, 0, alpha=0.8))  # 红色
            
            # 获取页面尺寸以确定右上角位置
            page_width = float(first_page.mediabox.width)
            page_height = float(first_page.mediabox.height)
            
            # 在右上角添加分数，只显示数字
            score_x = page_width - 180  # 距离右边180点（增加空间以适应大字体）
            score_y = page_height - 120  # 距离顶部120点，更靠下的位置，确保显示完整
            score_canvas.drawString(score_x, score_y, f"{score}")
            
            score_canvas.save()
            
            score_packet.seek(0)
            score_pdf = PdfReader(score_packet)
            
            # 合并分数到第一页
            first_page.merge_page(score_pdf.pages[0])
        
        output_pdf.add_page(first_page)
        
        # 添加其余页面
        for page in input_pdf.pages[1:]:
            output_pdf.add_page(page)
        
        # 添加新的一页用于评语（旋转显示）
        comment_packet = io.BytesIO()
        comment_canvas = canvas.Canvas(comment_packet, pagesize=letter)
        
        # 使用支持中文的字体
        comment_canvas.setFont(CHINESE_FONT_NAME, 14)
            
        comment_canvas.setFillColor(Color(1, 0, 0, alpha=0.8))  # 红色
        
        # 添加标题
        comment_canvas.setFont(CHINESE_FONT_NAME, 16)  # 使用中文字体，减小标题字号
        comment_canvas.drawString(50, 750, "批阅评语")
        
        # 设置评语内容字体
        comment_canvas.setFont(CHINESE_FONT_NAME, 14)  # 使用中文字体
        
        # 分行显示评语，确保处理中文换行并控制宽度不超出页面
        def split_text_to_lines(text, max_width, canvas, font_name=CHINESE_FONT_NAME, font_size=14):
            """将文本分割为适合页面宽度的行，支持中英文混合文本
            
            Args:
                text: 要分割的文本
                max_width: 最大允许宽度（点）
                canvas: Canvas对象，用于计算文本宽度
                font_name: 字体名称
                font_size: 字体大小
                
            Returns:
                分割后的行列表
            """
            lines = []
            if not text:
                return lines
            
            current_line = ""
            for char in text:
                # 如果是换行符，直接结束当前行
                if char == '\n':
                    lines.append(current_line)
                    current_line = ""
                    continue
                
                # 尝试将当前字符添加到当前行
                test_line = current_line + char
                test_width = canvas.stringWidth(test_line, font_name, font_size)
                
                # 如果添加后超过最大宽度
                if test_width > max_width:
                    # 如果当前行不为空，添加到结果列表
                    if current_line:
                        lines.append(current_line)
                        current_line = char
                    else:
                        # 如果当前行为空（单个字符就超过宽度），直接添加
                        lines.append(char)
                else:
                    # 添加当前字符到当前行
                    current_line = test_line
            
            # 添加最后一行
            if current_line:
                lines.append(current_line)
            
            return lines
        
        # 计算页面可用于评语的最大宽度（页面宽度减去左右边距）
        # 使用letter尺寸的宽度，约612点，左右各留50点边距
        max_comment_width = 612 - 100  # 最大宽度为512点
        
        lines = []
        for original_line in comments.split('\n'):
            if not original_line.strip():
                lines.append('')
            else:
                # 分割当前行
                line_parts = split_text_to_lines(original_line, max_comment_width, comment_canvas)
                lines.extend(line_parts)
        
        y_pos = 690  # 调整评语开始位置，因为添加了日期
        for line in lines:
            if y_pos < 100:  # 如果空间不够，开始新页
                    comment_canvas.showPage()
                    comment_canvas.setFont(CHINESE_FONT_NAME, 14)  # 确保新页面也使用中文字体和14点字号
                    comment_canvas.setFillColor(Color(1, 0, 0, alpha=0.8))
                    y_pos = 750
            comment_canvas.drawString(50, y_pos, line)
            y_pos -= 25  # 增大行距以适应更大的字号
        
        comment_canvas.save()
        
        comment_packet.seek(0)
        comment_pdf = PdfReader(comment_packet)
        
        # 添加评语页
        output_pdf.add_page(comment_pdf.pages[0])
        
        with open(output_path, "wb") as out_f:
            output_pdf.write(out_f)

    def add_annotations(self, file_path: str, annotations: List[Dict[str, Any]], output_path: str) -> bool:
        """在PDF上添加批注
        
        Args:
            file_path: 原始PDF文件路径
            annotations: 批注列表，每个批注包含页码、位置和内容
            output_path: 输出文件路径
            
        Returns:
            bool: 处理是否成功
        """
        try:
            input_pdf = PdfReader(file_path)
            output_pdf = PdfWriter()
            
            # 处理每一页
            for page_num in range(len(input_pdf.pages)):
                page = input_pdf.pages[page_num]
                
                # 找出当前页的所有批注
                page_annotations = [a for a in annotations if a.get("page", 0) == page_num]
                
                if page_annotations:
                    # 创建批注层
                    packet = io.BytesIO()
                    can = canvas.Canvas(packet, pagesize=letter)
                    can.setFont("Helvetica", 10)
                    
                    # 添加每个批注
                    for annotation in page_annotations:
                        x = annotation.get("x", 50)
                        y = annotation.get("y", 50)
                        text = annotation.get("text", "")
                        color = annotation.get("color", "red")
                        
                        # 设置颜色
                        if color == "red":
                            can.setFillColor(Color(1, 0, 0, alpha=0.7))
                        elif color == "green":
                            can.setFillColor(Color(0, 1, 0, alpha=0.7))
                        elif color == "blue":
                            can.setFillColor(Color(0, 0, 1, alpha=0.7))
                        else:
                            can.setFillColor(Color(0, 0, 0, alpha=0.7))
                        
                        # 绘制批注
                        can.drawString(x, y, text)
                    
                    can.save()
                    packet.seek(0)
                    annotation_pdf = PdfReader(packet)
                    
                    # 合并批注到当前页
                    page.merge_page(annotation_pdf.pages[0])
                
                output_pdf.add_page(page)
            
            # 保存输出文件
            with open(output_path, "wb") as out_f:
                output_pdf.write(out_f)
                
            return True
            
        except Exception as e:
            logger.error(f"添加批注时出错: {e}", exc_info=True)
            return False
    
    def add_checkmarks(self, file_path: str, output_path: str) -> None:
        """在PDF每页添加对号标注"""
        input_pdf = PdfReader(file_path)
        output_pdf = PdfWriter()
        num_pages = len(input_pdf.pages)

        for page_num in range(num_pages):
            page = input_pdf.pages[page_num]
            
            # 最后一页不加对号
            if page_num == num_pages - 1:
                output_pdf.add_page(page)
                continue
            
            # 获取页面尺寸
            page_width = float(page.mediabox.width)
            page_height = float(page.mediabox.height)
            
            # 计算页面中心位置
            center_x = page_width / 2
            center_y = page_height / 2

            # 创建对号标记
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=(page_width, page_height))
            can.setFont("Helvetica", 96)  # 更大的字号
            can.setFillColor(Color(1, 0, 0, alpha=0.8))  # 红色对号

            # 在每页中央添加对号
            # 调整对号位置以确保居中显示
            text_width = can.stringWidth("✓", "Helvetica", 96)
            x = center_x - text_width / 2
            y = center_y - 48  # 考虑字号高度的一半
            can.drawString(x, y, "✓")

            can.save()
            packet.seek(0)
            checkmark_pdf = PdfReader(packet)

            # 合并对号到当前页
            page.merge_page(checkmark_pdf.pages[0])
            output_pdf.add_page(page)

        with open(output_path, "wb") as out_f:
            output_pdf.write(out_f)


class MarkdownProcessor(DocumentProcessor):
    """Markdown文档处理实现"""

    def extract_text(self, file_path: str) -> str:
        """从Markdown中提取文本"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def add_comments_and_score(self, file_path: str, comments: str, score: int, output_path: str, add_score: bool = True) -> None:
        """在Markdown上添加评语和分数"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 添加分数和评语
        result = content
        if add_score:
            result += f"\n\n---\n## 评分结果\n**分数**: {score}\n\n---\n"
        result += f"## 教师评语\n{comments}\n"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(result)

    def add_checkmarks(self, file_path: str, output_path: str) -> None:
        """在Markdown文档的每个主要部分添加对号"""
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        result_lines = []
        for line in lines:
            # 在以#开头的标题行后添加对号
            if line.startswith('#'):
                result_lines.append(line)
                result_lines.append("✓\n")
            else:
                result_lines.append(line)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.writelines(result_lines)


class WordProcessor(DocumentProcessor):
    """Word文档处理实现"""
    
    def convert_to_pdf(self, word_path: str, pdf_path: str) -> bool:
        """将Word文档转换为PDF格式
        
        Args:
            word_path: Word文档路径
            pdf_path: 输出PDF文件路径
            
        Returns:
            bool: 转换是否成功
        """
        # 首先尝试使用win32com（适用于Windows环境）
        if WIN32COM_AVAILABLE:
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                # 打开文档
                doc = word.Documents.Open(word_path)
                
                # 转换为PDF
                # 17表示PDF格式
                doc.SaveAs(pdf_path, FileFormat=17)
                
                doc.Close()
                word.Quit()
                
                logger.info(f"成功将 {word_path} 转换为 PDF (使用win32com)")
                return True
                
            except Exception as e:
                logger.error(f"使用win32com转换Word到PDF时出错: {e}", exc_info=True)
                try:
                    word.Quit()
                except:
                    pass
        
        # 如果win32com不可用或失败，尝试使用LibreOffice命令行工具（适用于Linux环境）
        logger.info("尝试使用LibreOffice转换Word到PDF...")
        import subprocess
        
        try:
            # 构建LibreOffice命令
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf:writer_pdf_Export",
                "--outdir",
                os.path.dirname(pdf_path),
                word_path
            ]
            
            # 执行命令
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode == 0:
                # 当使用--outdir参数时，LibreOffice会将生成的PDF文件放在指定的输出目录
                # 生成的PDF文件名与源文件名相同，但扩展名为.pdf
                word_filename = os.path.basename(word_path)
                generated_pdf = os.path.join(
                    os.path.dirname(pdf_path),
                    os.path.splitext(word_filename)[0] + ".pdf"
                )
                
                if os.path.exists(generated_pdf):
                    # 如果目标路径与生成路径不同，移动文件
                    if generated_pdf != pdf_path:
                        os.rename(generated_pdf, pdf_path)
                    logger.info(f"成功将 {word_path} 转换为 PDF (使用LibreOffice)")
                    return True
                else:
                    logger.error(f"LibreOffice命令执行成功，但未生成输出文件: {generated_pdf}")
                    # 打印调试信息
                    logger.debug(f"LibreOffice输出: {result.stdout}")
                    logger.debug(f"LibreOffice错误: {result.stderr}")
                    return False
            else:
                logger.error(f"LibreOffice命令执行失败: {result.stderr}")
                return False
                
        except Exception as e:
            logger.error(f"使用LibreOffice转换Word到PDF时出错: {e}", exc_info=True)
            return False

    def extract_text(self, file_path: str) -> str:
        """从Word文档中提取文本"""
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        else:  # .doc
            if not WIN32COM_AVAILABLE:
                logger.warning("win32com not available, cannot process .doc files. Please convert to .docx format.")
                return "Error: Cannot process .doc files in this environment. Please convert to .docx format."
            
            word = win32com.client.Dispatch("Word.Application")
            doc = word.Documents.Open(file_path)
            text = doc.Content.Text
            doc.Close()
            word.Quit()
            return text

    def add_comments_and_score(self, file_path: str, comments: str, score: int, output_path: str, add_score: bool = True) -> None:
        """在Word文档上添加评语和分数
        
        Args:
            file_path: 输入Word文档路径
            comments: 要添加的评语
            score: 要添加的分数
            output_path: 输出Word文档路径
            add_score: 是否添加分数，默认为True
        """
        if WIN32COM_AVAILABLE and file_path.lower().endswith('.doc'):
            # 使用win32com处理.doc文件
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            
            try:
                # 只在需要时添加分数
                if add_score:
                    # 在首页右上角添加分数
                    selection = word.Selection
                    
                    # 转到首页
                    selection.GoTo(What=1, Which=1, Count=1)  # 1=wdGoToPage, 1=wdGoToFirst
                    
                    # 转到页首
                    selection.HomeKey(Unit=6)  # 6=wdStory
                    
                    # 设置为右对齐
                    selection.ParagraphFormat.Alignment = 2  # 2=wdAlignParagraphRight
                    
                    # 插入分数，只显示数字
                    selection.TypeText(f"{score}")
                    
                    # 设置分数格式
                    selection.Font.Bold = True
                    selection.Font.Size = 48  # 字号增加到48pt
                    selection.Font.Color = 16711680  # 红色
                
                # 添加分页符
                selection.TypeParagraph()
                selection.InsertBreak(7)  # 7=wdPageBreak
                
                # 添加评语标题
                selection.TypeText("批阅评语")
                selection.Font.Bold = True
                selection.Font.Size = 24
                selection.Font.Color = 16711680  # 红色
                selection.ParagraphFormat.Alignment = 1  # 1=wdAlignParagraphCenter
                
                # 添加新段落
                selection.TypeParagraph()
                selection.ParagraphFormat.Alignment = 0  # 0=wdAlignParagraphLeft
                selection.Font.Bold = False
                selection.Font.Size = 14  # 进一步减小字号
                selection.Font.Color = 0  # 黑色
                
                # 添加评语内容，确保正确处理中文
                selection.TypeText(comments)
                
                # 保存文档
                doc.SaveAs(output_path)
            finally:
                doc.Close()
                word.Quit()
        else:
            # 使用python-docx处理.docx文件
            doc = Document(file_path)
            
            # 只在需要时添加分数
            if add_score:
                # 在首页右上角添加分数，只显示数字
                score_para = doc.add_paragraph()
                score_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
                # 调整段落间距，为大字体提供足够空间
                from docx.shared import Pt
                score_para.paragraph_format.space_before = Pt(10)
                score_para.paragraph_format.space_after = Pt(10)
                
                score_run = score_para.add_run(f"{score}")
                score_run.bold = True
                score_run.font.size = 48  # 字号增加到48pt
            
            # 添加新的一页用于评语
            doc.add_page_break()
            
            # 添加标题
            title_para = doc.add_paragraph()
            title_run = title_para.add_run("批阅评语")
            title_run.bold = True
            title_run.font.size = 24  # 更大的字号 (24pt)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加评语内容，确保正确处理中文换行
            comments_para = doc.add_paragraph()
            comments_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # 设置字体以支持中文显示
            from docx.shared import RGBColor, Pt
            
            # 分段落添加评语
            comment_lines = comments.split('\n')
            for line in comment_lines:
                if line.strip():
                    run = comments_para.add_run(line.strip())
                    run.font.name = 'Arial Unicode MS'  # 使用支持中文的字体
                    run.font.size = Pt(14)  # 进一步减小字号到14pt
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 确保文本颜色为黑色，避免红色方块
                    comments_para.add_run('\n')
            
            doc.save(output_path)

    def add_checkmarks(self, file_path: str, output_path: str) -> None:
        """在Word文档每页添加对号标注"""
        if WIN32COM_AVAILABLE and file_path.lower().endswith('.doc'):
            # 使用win32com处理.doc文件
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            
            try:
                # 获取所有页面
                pages = doc.ComputeStatistics(2)  # 2表示wdStatisticPages
                
                # 在每页中央添加对号（最后一页除外）
                for page_num in range(1, pages):  # 只处理到倒数第二页
                    # 转到页尾
                    word.Selection.GoTo(What=1, Which=2, Count=page_num)  # 1=wdGoToPage, 2=wdGoToLast
                    
                    # 插入对号
                    selection = word.Selection
                    selection.TypeText("✓")
                    
                    # 设置对号格式
                    selection.Font.Size = 96  # 更大的字号
                    selection.Font.Color = 16711680  # 红色 (RGB: 255, 0, 0)
                    
                    # 居中对齐
                    selection.ParagraphFormat.Alignment = 1  # 1=wdAlignParagraphCenter
                    
                    # 调整位置到页面中央（近似）
                    selection.ParagraphFormat.SpaceBefore = 360  # 增加段前间距
                    selection.ParagraphFormat.SpaceAfter = 360  # 增加段后间距
                    
                    # 添加分页符
                    selection.TypeParagraph()
                    selection.InsertBreak(7)  # 7=wdPageBreak
                
                # 保存文档
                doc.SaveAs(output_path)
            finally:
                doc.Close()
                word.Quit()
        else:
            # 使用python-docx处理.docx文件
            doc = Document(file_path)
            
            # 在文档开头添加分数和评语后，在每页添加对号（最后一页除外）
            # 注意：python-docx不直接支持分页信息，这里采用在段落间添加的方式
            paragraphs = doc.paragraphs
            
            # 计算需要添加对号的数量，留出最后一段不添加对号
            max_checkmarks = len(paragraphs) // 5 - 1 if len(paragraphs) > 5 else 0
            checkmark_count = 0
            
            # 在每个主要段落后添加对号
            for i, para in enumerate(paragraphs):
                if i > 0 and i % 5 == 0 and checkmark_count < max_checkmarks:  # 每5个段落添加一个对号（近似每页一个），留出最后一页
                    checkmark_para = doc.add_paragraph()
                    checkmark_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    checkmark_run = checkmark_para.add_run("✓")
                    checkmark_run.font.size = 96  # 更大的字号
                    from docx.shared import RGBColor
                    checkmark_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
                    
                    # 调整间距
                    checkmark_para.paragraph_format.space_before = 1000000  # 增加段前间距
                    checkmark_para.paragraph_format.space_after = 1000000  # 增加段后间距
                    checkmark_count += 1
            
            doc.save(output_path)    