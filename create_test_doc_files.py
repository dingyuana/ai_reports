#!/usr/bin/env python3
"""
文件名: create_test_doc_files.py
作用: 创建测试用的.doc和.docx文件
"""

import os
import subprocess
import logging

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_docx_file(docx_path):
    """创建一个简单的测试.docx文件"""
    try:
        from docx import Document
        
        logger.info(f"创建测试.docx文件: {docx_path}")
        
        # 创建新的Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading('Test Student Report', 0)
        
        # 添加段落
        doc.add_paragraph('This is a test report content.')
        doc.add_paragraph('This report contains the following sections:')
        
        # 添加一些测试内容
        sections = [
            '1. Introduction',
            'In this section, we introduce the project background and objectives.',
            
            '2. Methods',
            'This section describes our research methods and tools.',
            
            '3. Results',
            'Here we present our research findings and data analysis results.',
            
            '4. Discussion',
            'This section discusses the implications and potential impact of our findings.',
            
            '5. Conclusion',
            'Finally, we summarize the main findings and recommendations of the study.'
        ]
        
        for section in sections:
            doc.add_paragraph(section)
        
        # 保存文档
        doc.save(docx_path)
        logger.info(f"成功创建.docx文件: {docx_path}")
        return True
        
    except Exception as e:
        logger.error(f"创建测试.docx文件失败: {e}", exc_info=True)
        return False

def convert_docx_to_doc(docx_path, doc_path):
    """使用LibreOffice将.docx文件转换为.doc文件"""
    try:
        logger.info(f"将.docx文件转换为.doc文件: {docx_path} -> {doc_path}")
        
        # 构建LibreOffice命令
        cmd = [
            "libreoffice",
            "--headless",
            "--convert-to",
            "doc:MS Word 97",
            "--outdir",
            os.path.dirname(doc_path),
            docx_path
        ]
        
        # 执行命令
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            # LibreOffice会生成一个与源文件同名的.doc文件
            generated_doc = os.path.join(
                os.path.dirname(doc_path),
                os.path.splitext(os.path.basename(docx_path))[0] + ".doc"
            )
            
            if os.path.exists(generated_doc):
                if generated_doc != doc_path:
                    os.rename(generated_doc, doc_path)
                logger.info(f"成功将.docx文件转换为.doc文件: {doc_path}")
                return True
            else:
                logger.error(f"LibreOffice命令执行成功，但未生成.doc文件")
                logger.debug(f"LibreOffice输出: {result.stdout}")
                logger.debug(f"LibreOffice错误: {result.stderr}")
                return False
        else:
            logger.error(f"LibreOffice命令执行失败: {result.stderr}")
            return False
            
    except Exception as e:
        logger.error(f"转换.docx到.doc文件失败: {e}", exc_info=True)
        return False

def main():
    """主函数"""
    # 创建测试目录
    test_dir = os.path.join(os.getcwd(), "student_reports")
    os.makedirs(test_dir, exist_ok=True)
    
    # 测试文件路径
    docx_path = os.path.join(test_dir, "test_report.docx")
    doc_path = os.path.join(test_dir, "test_report.doc")
    
    try:
        # 创建测试.docx文件
        if not create_docx_file(docx_path):
            logger.error("创建.docx文件失败")
            return 1
        
        # 将.docx文件转换为.doc文件
        if not convert_docx_to_doc(docx_path, doc_path):
            logger.error("转换.docx到.doc文件失败")
            return 1
        
        logger.info("所有测试文件创建成功！")
        return 0
        
    except Exception as e:
        logger.error(f"测试文件创建过程中出现意外错误: {e}", exc_info=True)
        return 1

if __name__ == "__main__":
    exit(main())